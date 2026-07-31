from docx import Document
doc = Document()
doc.add_page_break()
doc.add_heading("4.3 Traceability Matrix", level=2)
paragraph = doc.add_paragraph(
    "The traceability matrix maps each User Acceptance Test Case to the "
    "corresponding Functional Requirement and UAT Scenario."
)
paragraph.paragraph_format.line_spacing = 1.5
headers = [
"Test Case ID",
"Functional Requirement",
"UAT Scenario",
"Requirement Covered"
]
table1 = doc.add_table(rows=1, cols=len(headers))
table1.style = "Table Grid"
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h
matrix = [
["TC-001","FR-01","UAT-001","Payroll Officer Login"],
["TC-002","FR-01","UAT-001 ","Reject Invalid Login"],
["TC-003","FR-01","UAT-002","Initiate Payroll"],
["TC-004","FR-03","UAT-003","Prevent Duplicate Payroll"],
["TC-005","FR-04","UAT-004","Submit Payroll"],
["TC-006","FR-04","UAT-005","Approve Payroll"],
["TC-007","FR-05","UAT-006","Reject Payroll"],
["TC-008","FR-04","UAT-007","Employee Views Payslip"],
["TC-009","FR-10","UAT-008","Update PAYE Rates"],
["TC-010","FR-09","Support","Loan Deduction"],
["TC-011","FR-02","Support","Missing Salary Structure"],
["TC-012","FR-06","Support","Missing PAYE Configuration"],
["TC-013","FR-04","Support","Generate Payroll Report"],
["TC-014","FR-01","Support","Employee Search"],
["TC-015","FR-10","Support","Unauthorized Access"]
]
for item in matrix:
    row = table1.add_row().cells
    for i in range(len(item)):
        row[i].text = item[i]
doc.save("UAT.docx")
print("DOCUMENT IS SUCCESSFULLY CREATED.")